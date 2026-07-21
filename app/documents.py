"""Парсинг прикреплённых документов (§16).

Белый список форматов, проверка расширения и MIME, парсинг во временном
каталоге, никакого исполнения содержимого. Возвращает извлечённый текст и/или
изображения (для vision-пути) в виде data-URL.

ИБ:
- docx — защита от zip-бомб (лимит суммарного распакованного размера);
- изображения перекодируются через Pillow (отсечение метаданных/полиглотов);
- PDF: приоритет — текстовый слой (pypdf), дешевле по токенам; скан —
  растеризация pdftoppm и передача модели как изображения (vision, mmproj).

Изображения и сканы распознаёт сама мультимодальная модель (llama.cpp с mmproj).
"""
from __future__ import annotations

import base64
import csv
import io
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass, field
from pathlib import Path

from app.config import settings

# Белый список: расширение -> допустимые MIME (пустой набор — MIME не проверяем строго)
ALLOWED: dict[str, set[str]] = {
    ".txt": {"text/plain"},
    ".md": {"text/plain", "text/markdown", "application/octet-stream"},
    ".csv": {"text/csv", "text/plain", "application/vnd.ms-excel"},
    ".docx": {"application/vnd.openxmlformats-officedocument.wordprocessingml.document",
              "application/zip", "application/octet-stream"},
    ".xlsx": {"application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
              "application/zip", "application/octet-stream"},
    ".pdf": {"application/pdf", "application/octet-stream"},
    ".png": {"image/png"},
    ".jpg": {"image/jpeg"},
    ".jpeg": {"image/jpeg"},
}

XLSX_MAX_ROWS = 200          # ограничение строк на лист (§16)
DOCX_MAX_UNCOMPRESSED = 200 * 1024 * 1024  # защита от zip-бомб: 200 МБ распакованного
CHARS_PER_TOKEN = 3          # грубая оценка длины (§16)
IMAGE_TOKEN_BUDGET = 800     # ориентировочный бюджет токенов на изображение
IMAGE_MAX_DIM = 2000         # макс. сторона изображения перед отправкой модели


class DocumentError(Exception):
    """Ошибка обработки документа — понятный текст для пользователя."""


@dataclass
class ParsedDocument:
    filename: str
    text: str = ""
    images: list[str] = field(default_factory=list)  # data-URL (base64)
    warnings: list[str] = field(default_factory=list)

    @property
    def token_estimate(self) -> int:
        return len(self.text) // CHARS_PER_TOKEN + len(self.images) * IMAGE_TOKEN_BUDGET


def _ext(filename: str) -> str:
    return Path(filename).suffix.lower()


def validate(filename: str, content_type: str, size: int) -> str:
    ext = _ext(filename)
    if ext not in ALLOWED:
        raise DocumentError(f"Формат {ext or '(без расширения)'} не поддерживается")
    limit = settings.max_upload_mb * 1024 * 1024
    if size > limit:
        raise DocumentError(f"Файл больше {settings.max_upload_mb} МБ")
    mime = (content_type or "").split(";")[0].strip().lower()
    allowed_mimes = ALLOWED[ext]
    if mime and mime not in allowed_mimes:
        raise DocumentError(f"MIME-тип {mime} не соответствует расширению {ext}")
    return ext


# --- Текстовые форматы ---

def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8", "cp1251", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", "replace")


def _csv_to_markdown(data: bytes) -> str:
    text = _decode_text(data)
    sample = text[:2048]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t")
    except csv.Error:
        dialect = csv.excel
    rows = list(csv.reader(io.StringIO(text), dialect))
    rows = [r for r in rows if any(cell.strip() for cell in r)]
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]

    def fmt(cells: list[str]) -> str:
        return "| " + " | ".join(c.replace("|", "\\|").strip() for c in cells) + " |"

    lines = [fmt(rows[0]), "| " + " | ".join(["---"] * width) + " |"]
    lines += [fmt(r) for r in rows[1:]]
    return "\n".join(lines)


def _parse_docx(data: bytes) -> str:
    import docx  # python-docx

    # Защита от zip-бомб: проверяем суммарный распакованный размер
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        total = sum(info.file_size for info in zf.infolist())
        if total > DOCX_MAX_UNCOMPRESSED:
            raise DocumentError("Документ отклонён: слишком большой распакованный размер")

    document = docx.Document(io.BytesIO(data))
    parts: list[str] = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in document.tables:
        rows = []
        for row in table.rows:
            rows.append("| " + " | ".join(c.text.replace("|", "\\|").strip()
                                          for c in row.cells) + " |")
        if rows:
            width = table.rows[0].cells and len(table.rows[0].cells) or 1
            rows.insert(1, "| " + " | ".join(["---"] * width) + " |")
            parts.append("\n".join(rows))
    return "\n\n".join(parts)


def _parse_xlsx(data: bytes) -> tuple[str, list[str]]:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    warnings: list[str] = []
    blocks: list[str] = []
    for ws in wb.worksheets:
        rows_out: list[list[str]] = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= XLSX_MAX_ROWS:
                warnings.append(f"Лист «{ws.title}»: показаны первые {XLSX_MAX_ROWS} строк")
                break
            cells = ["" if v is None else str(v) for v in row]
            if any(c.strip() for c in cells):
                rows_out.append(cells)
        if not rows_out:
            continue
        width = max(len(r) for r in rows_out)
        rows_out = [r + [""] * (width - len(r)) for r in rows_out]
        md = ["| " + " | ".join(c.replace("|", "\\|").strip() for c in rows_out[0]) + " |",
              "| " + " | ".join(["---"] * width) + " |"]
        md += ["| " + " | ".join(c.replace("|", "\\|").strip() for c in r) + " |"
               for r in rows_out[1:]]
        blocks.append(f"### Лист: {ws.title}\n\n" + "\n".join(md))
    wb.close()
    return "\n\n".join(blocks), warnings


# --- Изображения ---

def _image_to_data_url(data: bytes) -> str:
    """Перекодировать изображение через Pillow (санитайзинг) в PNG data-URL."""
    from PIL import Image

    try:
        with Image.open(io.BytesIO(data)) as img:
            img.load()
            rgb = img.convert("RGB")
            if max(rgb.size) > IMAGE_MAX_DIM:
                ratio = IMAGE_MAX_DIM / max(rgb.size)
                rgb = rgb.resize((int(rgb.width * ratio), int(rgb.height * ratio)))
            out = io.BytesIO()
            rgb.save(out, format="PNG")
    except Exception as exc:  # noqa: BLE001 — любой сбой декодера = отказ
        raise DocumentError(f"Не удалось обработать изображение: {exc}") from exc
    b64 = base64.b64encode(out.getvalue()).decode("ascii")
    return f"data:image/png;base64,{b64}"


# --- PDF ---

def _pdf_text(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise DocumentError(f"Не удалось прочитать PDF: {exc}") from exc
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 — битая страница не должна ронять всё
            continue
    return "\n\n".join(p for p in parts if p.strip()).strip()


def _pdf_rasterize(data: bytes, max_pages: int) -> list[bytes]:
    """Растеризовать страницы PDF в PNG через pdftoppm.
    max_pages > 0 — первые N страниц; max_pages <= 0 — все страницы."""
    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "in.pdf"
        src.write_bytes(data)
        cmd = ["pdftoppm", "-png", "-r", "150"]
        if max_pages and max_pages > 0:
            cmd += ["-l", str(max_pages)]
        cmd += [str(src), str(Path(tmp) / "page")]
        try:
            subprocess.run(cmd, check=True, capture_output=True, timeout=300)
        except FileNotFoundError:
            raise DocumentError(
                "Растеризация PDF недоступна: не установлен poppler-utils (pdftoppm)")
        except subprocess.CalledProcessError as exc:
            raise DocumentError(f"Ошибка растеризации PDF: {exc.stderr[:200]!r}")
        except subprocess.TimeoutExpired:
            raise DocumentError("Растеризация PDF превысила лимит времени")
        pages = sorted(Path(tmp).glob("page*.png"))
        return [p.read_bytes() for p in pages]


def _parse_pdf(data: bytes, doc: ParsedDocument, mode: str = "vision") -> None:
    """Режимы:
    - 'vision' (по умолчанию): страницы → картинки → mmproj (модель «видит» лист:
       вёрстку, таблицы, формулы, чертежи, сканы);
    - 'text': только текстовый слой (pypdf), дёшево; если слоя нет — ошибка;
    - 'auto': текст, если он есть, иначе vision (прежнее поведение).
    """
    if mode in ("text", "auto"):
        text = _pdf_text(data)
        if text:
            doc.text = text
            return
        if mode == "text":
            raise DocumentError(
                "В PDF нет текстового слоя — выберите режим «как картинку»")
        # auto: текста нет → падаем в vision ниже

    # vision (по умолчанию) или auto без текста — распознаёт мультимодальная модель.
    # Каждая страница — отдельная полная картинка. По умолчанию берём ВСЕ страницы
    # (VISION_MAX_PAGES=0); положительный лимит — предохранитель от переполнения
    # контекста (каждая страница = картинка на сотни-тысячи токенов).
    max_pages = settings.vision_max_pages
    if max_pages and max_pages > 0:
        images = _pdf_rasterize(data, max_pages + 1)
        if len(images) > max_pages:
            images = images[:max_pages]
            doc.warnings.append(
                f"PDF обрезан до {max_pages} страниц (VISION_MAX_PAGES), "
                "остальное не обработано")
    else:
        images = _pdf_rasterize(data, 0)  # все страницы
    if not images:
        raise DocumentError("PDF не содержит страниц для распознавания")
    if len(images) >= 15:
        doc.warnings.append(
            f"{len(images)} страниц отправлены картинками — это заметно заполнит "
            "контекст модели (следите за процентом контекста)")
    doc.images = [_image_to_data_url(img) for img in images]


# --- Точка входа ---

def parse_upload(filename: str, content_type: str, data: bytes,
                 pdf_mode: str = "vision") -> ParsedDocument:
    ext = validate(filename, content_type, len(data))
    doc = ParsedDocument(filename=filename)

    if ext in (".txt", ".md"):
        doc.text = _decode_text(data)
    elif ext == ".csv":
        doc.text = _csv_to_markdown(data)
    elif ext == ".docx":
        doc.text = _parse_docx(data)
    elif ext == ".xlsx":
        doc.text, doc.warnings = _parse_xlsx(data)
    elif ext in (".png", ".jpg", ".jpeg"):
        doc.images = [_image_to_data_url(data)]
    elif ext == ".pdf":
        _parse_pdf(data, doc, pdf_mode)

    if not doc.text.strip() and not doc.images:
        raise DocumentError("Из файла не извлечено ни текста, ни изображений")
    return doc
