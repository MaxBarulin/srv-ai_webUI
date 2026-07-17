"""Document parsing tests (§16): форматы, ИБ, vision-путь, интеграция с чатом."""
from __future__ import annotations

import base64
import io
import json
import sqlite3
import zipfile
from dataclasses import replace

import httpx
import pytest

from app import documents
from app import llm as llm_module
from app.config import settings
from app.documents import DocumentError, parse_upload
from tests.conftest import login_as
from tests.mock_llm import app as mock_llm_app

PASS = "doc-user-pass-01"


@pytest.fixture(autouse=True)
def mock_llm(monkeypatch):
    monkeypatch.setattr(llm_module, "_transport", httpx.ASGITransport(app=mock_llm_app))


@pytest.fixture()
def doc_user(client, make_user):
    make_user("doc-user", PASS)
    login_as(client, "doc-user", PASS)
    yield
    conn = sqlite3.connect(settings.db_path)
    try:
        conn.execute("DELETE FROM messages")
        conn.execute("DELETE FROM chats")
        conn.commit()
    finally:
        conn.close()


# --- Валидация ---

def test_reject_unknown_extension():
    with pytest.raises(DocumentError, match="не поддерживается"):
        parse_upload("evil.exe", "application/octet-stream", b"MZ")


def test_reject_mime_mismatch():
    with pytest.raises(DocumentError, match="MIME"):
        parse_upload("note.txt", "application/x-msdownload", b"hi")


def test_reject_oversize(monkeypatch):
    monkeypatch.setattr(documents, "settings", replace(settings, max_upload_mb=1))
    with pytest.raises(DocumentError, match="больше"):
        parse_upload("big.txt", "text/plain", b"x" * (2 * 1024 * 1024))


# --- Текстовые форматы ---

def test_parse_txt():
    doc = parse_upload("note.txt", "text/plain", "Привет, мир".encode())
    assert doc.text == "Привет, мир"
    assert doc.images == []


def test_parse_csv_to_markdown():
    csv_bytes = "имя,возраст\nИван,30\nПётр,25\n".encode()
    doc = parse_upload("data.csv", "text/csv", csv_bytes)
    assert "| имя | возраст |" in doc.text
    assert "| --- | --- |" in doc.text
    assert "| Иван | 30 |" in doc.text


def test_parse_csv_semicolon_and_cp1251():
    csv_bytes = "имя;город\nИван;Москва\n".encode("cp1251")
    doc = parse_upload("data.csv", "text/csv", csv_bytes)
    assert "| имя | город |" in doc.text
    assert "Москва" in doc.text


def test_parse_docx():
    import docx as docxlib
    d = docxlib.Document()
    d.add_paragraph("Первый абзац")
    table = d.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "2"
    buf = io.BytesIO()
    d.save(buf)
    doc = parse_upload("doc.docx", "", buf.getvalue())
    assert "Первый абзац" in doc.text
    assert "| A | B |" in doc.text
    assert "| 1 | 2 |" in doc.text


def test_parse_xlsx_with_row_limit(monkeypatch):
    import openpyxl
    monkeypatch.setattr(documents, "XLSX_MAX_ROWS", 5)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Данные"
    for i in range(20):
        ws.append([f"строка{i}", i])
    buf = io.BytesIO()
    wb.save(buf)
    doc = parse_upload("book.xlsx", "", buf.getvalue())
    assert "### Лист: Данные" in doc.text
    assert any("первые 5 строк" in w for w in doc.warnings)


def test_docx_zip_bomb_guard(monkeypatch):
    monkeypatch.setattr(documents, "DOCX_MAX_UNCOMPRESSED", 1000)
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("word/document.xml", b"0" * 100000)  # хорошо сжимается
    with pytest.raises(DocumentError, match="распакованный размер"):
        parse_upload("bomb.docx", "", buf.getvalue())


# --- Изображения ---

def _png_bytes(color=(255, 0, 0), size=(10, 10)) -> bytes:
    from PIL import Image
    buf = io.BytesIO()
    Image.new("RGB", size, color).save(buf, format="PNG")
    return buf.getvalue()


def test_parse_image_vision_path():
    doc = parse_upload("photo.png", "image/png", _png_bytes())
    assert len(doc.images) == 1
    assert doc.images[0].startswith("data:image/png;base64,")
    assert doc.text == ""
    # data-URL декодируется в валидный PNG (перекодировано через Pillow)
    raw = base64.b64decode(doc.images[0].split(",", 1)[1])
    assert raw[:8] == b"\x89PNG\r\n\x1a\n"


def test_reject_corrupt_image():
    with pytest.raises(DocumentError, match="изображение"):
        parse_upload("bad.png", "image/png", b"not really a png")


def test_large_image_downscaled(monkeypatch):
    monkeypatch.setattr(documents, "IMAGE_MAX_DIM", 100)
    from PIL import Image
    doc = parse_upload("big.png", "image/png", _png_bytes(size=(500, 300)))
    raw = base64.b64decode(doc.images[0].split(",", 1)[1])
    with Image.open(io.BytesIO(raw)) as img:
        assert max(img.size) <= 100


# --- PDF ---

def _pdf_with_text(text: str) -> bytes:
    # Минимальный PDF с текстовым слоем через pypdf/reportlab отсутствует;
    # используем pypdf для сборки страницы с текстом невозможно без writer content.
    # Упростим: соберём PDF через PIL (изображение) для скан-пути, а текстовый
    # путь проверим на реальном PDF из pypdf writer.
    from pypdf import PdfWriter
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def test_pdf_scan_vision_path(monkeypatch):
    # PDF без текстового слоя → растеризация → изображения (vision)
    monkeypatch.setattr(documents, "settings",
                        replace(settings, vision_max_pages=10))
    monkeypatch.setattr(documents, "_pdf_text", lambda data: "")
    monkeypatch.setattr(documents, "_pdf_rasterize",
                        lambda data, max_pages: [_png_bytes(), _png_bytes()])
    doc = parse_upload("scan.pdf", "application/pdf", b"%PDF-1.4 fake")
    assert len(doc.images) == 2
    assert all(u.startswith("data:image/png") for u in doc.images)


def test_pdf_scan_page_limit(monkeypatch):
    monkeypatch.setattr(documents, "settings",
                        replace(settings, vision_max_pages=2))
    monkeypatch.setattr(documents, "_pdf_text", lambda data: "")
    monkeypatch.setattr(documents, "_pdf_rasterize",
                        lambda data, max_pages: [_png_bytes() for _ in range(5)])
    doc = parse_upload("scan.pdf", "application/pdf", b"%PDF fake")
    assert len(doc.images) == 2
    assert any("обрезан" in w for w in doc.warnings)


def test_pdf_text_layer_preferred(monkeypatch):
    monkeypatch.setattr(documents, "_pdf_text", lambda data: "Извлечённый текст PDF")
    called = {"raster": False}
    monkeypatch.setattr(documents, "_pdf_rasterize",
                        lambda *a: called.__setitem__("raster", True) or [])
    doc = parse_upload("doc.pdf", "application/pdf", b"%PDF fake")
    assert doc.text == "Извлечённый текст PDF"
    assert doc.images == []
    assert called["raster"] is False  # растеризация не вызывалась




# --- Endpoint загрузки ---

def test_upload_endpoint(client, doc_user):
    files = {"file": ("note.txt", b"Test content", "text/plain")}
    r = client.post("/api/attachments", files=files)
    assert r.status_code == 200
    body = r.json()
    assert body["text"] == "Test content"
    assert body["images"] == []
    assert body["token_estimate"] >= 0


def test_upload_rejects_bad_type(client, doc_user):
    files = {"file": ("evil.exe", b"MZ", "application/octet-stream")}
    r = client.post("/api/attachments", files=files)
    assert r.status_code == 400


def test_upload_requires_auth(client):
    files = {"file": ("note.txt", b"hi", "text/plain")}
    assert client.post("/api/attachments", files=files).status_code == 401


# --- Интеграция вложений с чатом ---

def _parse_sse(text: str):
    from tests.test_chat import _parse_sse as p
    return p(text)


def test_chat_message_with_document_text(client, doc_user, monkeypatch):
    captured = []
    orig = llm_module.stream_chat

    def spy(messages, tools=None, **kwargs):
        captured.append(messages)
        return orig(messages, tools=tools)

    monkeypatch.setattr("app.routers.chat.stream_chat", spy)

    chat_id = client.post("/api/chats", json={}).json()["id"]
    r = client.post(f"/api/chats/{chat_id}/messages", json={
        "content": "Кратко о чём документ?",
        "use_tools": False,
        "attachments": [{"filename": "spec.txt", "text": "Технические требования к сварке."}],
    })
    assert r.status_code == 200
    user_msg = captured[0][-1]
    assert "Технические требования к сварке" in user_msg["content"]
    assert "[Документ: spec.txt]" in user_msg["content"]

    # В истории текст документа хранится отдельным вложением (не «стеной» в сообщении)
    msgs = client.get(f"/api/chats/{chat_id}/messages").json()
    stored = [m for m in msgs if m["role"] == "user"][0]
    assert stored["content"] == "Кратко о чём документ?"
    assert stored["attachments"] == [
        {"filename": "spec.txt", "text": "Технические требования к сварке."}]

    # В последующем вопросе документ восстанавливается в историю для LLM
    client.post(f"/api/chats/{chat_id}/messages",
                json={"content": "уточни", "use_tools": False})
    followup_history = captured[1]
    joined = " ".join(str(m["content"]) for m in followup_history)
    assert "[Документ: spec.txt]" in joined
    assert "Технические требования к сварке" in joined


def test_chat_message_with_image_multimodal(client, doc_user, monkeypatch):
    captured = []
    orig = llm_module.stream_chat

    def spy(messages, tools=None, **kwargs):
        captured.append(messages)
        return orig(messages, tools=tools)

    monkeypatch.setattr("app.routers.chat.stream_chat", spy)

    data_url = "data:image/png;base64," + base64.b64encode(_png_bytes()).decode()
    chat_id = client.post("/api/chats", json={}).json()["id"]
    r = client.post(f"/api/chats/{chat_id}/messages", json={
        "content": "Что на картинке?",
        "use_tools": False,
        "attachments": [{"filename": "photo.png", "images": [data_url]}],
    })
    assert r.status_code == 200

    user_msg = captured[0][-1]
    assert isinstance(user_msg["content"], list)
    kinds = [part["type"] for part in user_msg["content"]]
    assert "image_url" in kinds
    assert user_msg["content"][-1]["image_url"]["url"] == data_url

    # В БД изображение не сохраняется — только пометка-вложение с именем файла
    msgs = client.get(f"/api/chats/{chat_id}/messages").json()
    stored = [m for m in msgs if m["role"] == "user"][0]
    assert stored["attachments"] == [{"filename": "photo.png", "image": True}]
    # (ответ mock-LLM цитирует запрос, поэтому проверяем только сообщения пользователя)
    assert "base64" not in json.dumps([m for m in msgs if m["role"] == "user"])


def test_attachment_text_truncated(client, doc_user, monkeypatch):
    monkeypatch.setattr("app.routers.chat.MAX_ATTACHMENT_CHARS", 100)
    chat_id = client.post("/api/chats", json={}).json()["id"]
    r = client.post(f"/api/chats/{chat_id}/messages", json={
        "content": "вопрос",
        "use_tools": False,
        "attachments": [{"filename": "big.txt", "text": "А" * 500}],
    })
    events = _parse_sse(r.text)
    warnings = [d["detail"] for e, d in events if e == "doc_warning"]
    assert any("обрезан" in w for w in warnings)


def test_empty_message_without_attachments_rejected(client, doc_user):
    chat_id = client.post("/api/chats", json={}).json()["id"]
    r = client.post(f"/api/chats/{chat_id}/messages",
                    json={"content": "  ", "use_tools": False})
    assert r.status_code == 400
